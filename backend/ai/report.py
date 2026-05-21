"""
ai/report.py — Auto-generated PDF + Word reports
Sections: Executive Summary, Dataset Profile, Preprocessing Log, Model Card,
          Performance Metrics, SHAP Analysis, Fairness, Deployment Guide, Recommendations
"""
import os
import json
import logging
from datetime import datetime

import anthropic
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib import colors
from docx import Document

from models_db import TrainingJob, Model, Session as DBSession, Report

logger = logging.getLogger(__name__)
REPORTS_PATH = os.getenv("REPORTS_PATH", "./reports")
os.makedirs(REPORTS_PATH, exist_ok=True)

client = anthropic.AsyncAnthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))


async def generate_report(job_id: str, db) -> dict:
    job = db.query(TrainingJob).filter(TrainingJob.id == job_id).first()
    if not job:
        raise ValueError(f"Job {job_id} not found")

    session = db.query(DBSession).filter(DBSession.id == job.session_id).first()
    best_model = db.query(Model).filter(Model.id == job.best_model_id).first()
    all_models = db.query(Model).filter(Model.job_id == job_id).all()

    data = {
        "job_id": job_id,
        "model_type": job.model_type,
        "generated_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"),
        "domain": session.domain_mode if session else "general",
        "dataset_profile": session.profile_json if session else {},
        "best_model": {
            "algorithm": best_model.algorithm if best_model else "N/A",
            "model_type": best_model.model_type if best_model else "ml",
            "framework": best_model.framework if best_model else "sklearn",
            "metrics": best_model.metrics_json if best_model else {},
            "hyperparams": best_model.hyperparams_json if best_model else {},
            "shap": best_model.shap_json if best_model else {},
        },
        "all_models": [
            {"algorithm": m.algorithm, "model_type": m.model_type, "metrics": m.metrics_json}
            for m in all_models
        ],
    }

    narrative = await _generate_narrative(data)
    data["narrative"] = narrative

    pdf_path = os.path.join(REPORTS_PATH, f"report_{job_id}.pdf")
    _build_pdf(data, pdf_path)

    docx_path = os.path.join(REPORTS_PATH, f"report_{job_id}.docx")
    _build_docx(data, docx_path)

    report = Report(job_id=job_id, pdf_path=pdf_path, docx_path=docx_path)
    db.add(report)
    db.commit()

    logger.info(f"Report generated: {pdf_path}")
    return {"pdf": pdf_path, "docx": docx_path, "report_id": report.id}


async def _generate_narrative(data: dict) -> dict:
    prompt = f"""
You are writing a Unified AI Platform model report. Given the training results below,
write THREE sections:
1. EXECUTIVE_SUMMARY: 2-3 sentences for a business audience
2. METRICS_INTERPRETATION: Plain-English explanation of model performance
3. RECOMMENDATIONS: 3-5 bullet points for model improvement

Training Results:
Best Model: {data['best_model']['algorithm']} ({data['best_model']['model_type']}/{data['best_model']['framework']})
Metrics: {data['best_model']['metrics']}
Top SHAP Features: {data['best_model']['shap']}
Domain: {data['domain']}
All Models: {[m['algorithm'] for m in data['all_models']]}

Return ONLY valid JSON with keys: executive_summary, metrics_interpretation, recommendations
"""
    response = await client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=1000,
        messages=[{"role": "user", "content": prompt}],
    )
    try:
        text = response.content[0].text.strip().replace("```json", "").replace("```", "").strip()
        return json.loads(text)
    except Exception:
        return {
            "executive_summary": response.content[0].text[:300],
            "metrics_interpretation": "See metrics table below.",
            "recommendations": ["Review SHAP values for feature insights."],
        }


def _build_pdf(data: dict, path: str):
    doc = SimpleDocTemplate(path, pagesize=letter,
                            rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    story = []

    def h1(t): return Paragraph(t, ParagraphStyle("H1", parent=styles["Heading1"], fontSize=18, spaceAfter=12, textColor=colors.HexColor("#1a1a2e")))
    def h2(t): return Paragraph(t, ParagraphStyle("H2", parent=styles["Heading2"], fontSize=13, spaceAfter=8, textColor=colors.HexColor("#16213e")))
    def body(t): return Paragraph(t, styles["Normal"])

    story.append(h1("Unified AI Platform — Model Report"))
    story.append(body(f"Generated: {data['generated_at']} | Job ID: {data['job_id']} | Type: {data['model_type'].upper()}"))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
    story.append(Spacer(1, 0.2 * inch))

    narrative = data.get("narrative", {})
    story.append(h2("Executive Summary"))
    story.append(body(narrative.get("executive_summary", "")))
    story.append(Spacer(1, 0.2 * inch))

    best = data.get("best_model", {})
    story.append(h2("Best Model — Model Card"))
    model_data = [
        ["Attribute", "Value"],
        ["Algorithm", best.get("algorithm", "N/A")],
        ["Model Type", best.get("model_type", "ml").upper()],
        ["Framework", best.get("framework", "sklearn")],
        ["Domain", data.get("domain", "general").capitalize()],
    ]
    for k, v in (best.get("hyperparams") or {}).items():
        model_data.append([k, str(v)])
    t = Table(model_data, colWidths=[2.5 * inch, 3.5 * inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.2 * inch))

    story.append(h2("Performance Metrics"))
    story.append(body(narrative.get("metrics_interpretation", "")))
    metrics = best.get("metrics") or {}
    if metrics:
        metric_rows = [["Metric", "Score"]] + [[k, str(v)] for k, v in metrics.items()]
        mt = Table(metric_rows, colWidths=[3 * inch, 3 * inch])
        mt.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f3460")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.lightblue, colors.white]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("PADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(mt)
    story.append(Spacer(1, 0.2 * inch))

    story.append(h2("All Models Comparison"))
    if data.get("all_models"):
        all_rows = [["Algorithm", "Type", "Key Metric"]]
        for m in data["all_models"]:
            metrics_str = ", ".join(f"{k}={v}" for k, v in (m.get("metrics") or {}).items())[:60]
            all_rows.append([m.get("algorithm", ""), m.get("model_type", "").upper(), metrics_str])
        ct = Table(all_rows, colWidths=[2 * inch, 1.5 * inch, 3 * inch])
        ct.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#533483")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.lavender, colors.white]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("PADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(ct)
    story.append(Spacer(1, 0.2 * inch))

    top_features = (best.get("shap") or {}).get("top_features", [])
    if top_features:
        story.append(h2("Top Feature Importances (SHAP)"))
        shap_rows = [["Feature", "SHAP Importance"]] + \
                    [[f["feature"], str(f["importance"])] for f in top_features[:10]]
        st = Table(shap_rows, colWidths=[3.5 * inch, 2.5 * inch])
        st.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#533483")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.lavender, colors.white]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("PADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(st)
    story.append(Spacer(1, 0.2 * inch))

    story.append(h2("AI Recommendations"))
    recs = narrative.get("recommendations", [])
    for r in (recs if isinstance(recs, list) else [str(recs)]):
        story.append(body(f"• {r}"))

    story.append(Spacer(1, 0.3 * inch))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
    story.append(body("<i>Unified AI Platform — Confidential</i>"))
    doc.build(story)


def _build_docx(data: dict, path: str):
    doc = Document()
    doc.add_heading("Unified AI Platform — Model Report", 0)
    doc.add_paragraph(f"Generated: {data['generated_at']} | Job: {data['job_id']} | Type: {data['model_type'].upper()}")

    narrative = data.get("narrative", {})
    best = data.get("best_model", {})

    doc.add_heading("Executive Summary", 1)
    doc.add_paragraph(narrative.get("executive_summary", ""))

    doc.add_heading("Best Model", 1)
    doc.add_paragraph(f"Algorithm: {best.get('algorithm', 'N/A')} | Type: {best.get('model_type', 'ml').upper()} | Framework: {best.get('framework', 'sklearn')}")

    doc.add_heading("Performance Metrics", 1)
    doc.add_paragraph(narrative.get("metrics_interpretation", ""))
    metrics = best.get("metrics") or {}
    if metrics:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = "Metric", "Score"
        for k, v in metrics.items():
            row = table.add_row().cells
            row[0].text, row[1].text = k, str(v)

    doc.add_heading("Top SHAP Features", 1)
    for f in (best.get("shap") or {}).get("top_features", [])[:10]:
        doc.add_paragraph(f"• {f['feature']}: {f['importance']}", style="List Bullet")

    doc.add_heading("Recommendations", 1)
    recs = narrative.get("recommendations", [])
    for r in (recs if isinstance(recs, list) else [str(recs)]):
        doc.add_paragraph(f"• {r}", style="List Bullet")

    doc.save(path)
